"""玄策 · Step4 验证层
对生成的 docx 做质量检查（空白页/空段落/表格/编码）。
用法: python postcheck.py <docx路径>
"""
import sys
from pathlib import Path
from docx import Document


def check(path: Path) -> int:
    doc = Document(str(path))
    problems = []
    checks = 0

    def add_check(name: str, ok: bool, detail: str = ""):
        nonlocal checks
        checks += 1
        status = "PASS" if ok else "FAIL"
        print(f"  [{status}] {name}{(' — ' + detail) if detail and not ok else ''}")
        if not ok:
            problems.append(name)

    # 1. 文件存在且可打开
    add_check("文件可打开", True, path.name)

    # 2. 有正文内容
    text = "\n".join(p.text for p in doc.paragraphs)
    add_check("有正文内容", len(text.strip()) > 0, "文档为空")

    # 3. 有表格
    add_check("含概览表格", len(doc.tables) > 0, "缺少概览表格")

    # 4. 段落数合理
    n_paras = len(doc.paragraphs)
    add_check("段落数合理", n_paras >= 20, f"仅 {n_paras} 段")

    # 5. 无大量空白页（连续空段落过多）——封面留白属正常，阈值放宽到 60%
    empty_runs = sum(1 for p in doc.paragraphs if not p.text.strip())
    add_check("无异常空白", empty_runs < len(doc.paragraphs) * 0.6, f"{empty_runs} 空段")

    # 6. 有"玄机"关键词
    add_check("包含核心关键词", "玄机" in text, "缺少'玄机'")

    # 7. 无乱码（� 字符）
    add_check("无乱码字符", "�" not in text and "??" not in text, "检测到乱码")

    # 8. 编码正确
    try:
        text.encode("utf-8")
        add_check("UTF-8编码", True)
    except Exception as e:
        add_check("UTF-8编码", False, str(e))

    # 9. 标题层级存在
    has_sections = any("概览" in t or "详细" in t or "一、" in t or "二、" in t for t in [p.text for p in doc.paragraphs])
    add_check("章节结构完整", has_sections, "缺少章节标题")

    print(f"\n共 {checks} 项检查，{len(problems)} 项失败")
    return 1 if problems else 0


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python postcheck.py <docx路径>")
        sys.exit(1)
    sys.exit(check(Path(sys.argv[1])))
