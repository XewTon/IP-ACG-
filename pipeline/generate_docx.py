"""玄策 · Step3 生成层
读分析 JSON → python-docx → 玄机IP动态速报.docx
用法: python generate_docx.py [--date 2026-08-13] [--out 玄机IP动态速报.docx]
"""
import argparse
import json
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from config import ANALYZED_DIR, OUTPUT_DIR

# 玄策配色（墨+朱砂+宣纸）
INK = RGBColor(0x2A, 0x2E, 0x37)
RED = RGBColor(0xDA, 0x1E, 0x2B)
GOLD = RGBColor(0xD9, 0xA8, 0x45)
MUTED = RGBColor(0x6E, 0x68, 0x5C)

CAT_LABEL = {"company": "公司", "ipo": "IPO", "ip": "IP", "strategy": "战略", "industry": "行业"}
SCORE_LABEL = {5: "★★★★★ 重大", 4: "★★★★ 重要", 3: "★★★ 常规", 2: "★★ 边缘", 1: "★ 无关"}


def set_cn_font(run, name="微软雅黑", size=10, color=INK, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold


def build_cover(doc: Document, title: str, sub: str):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_cn_font(r, "微软雅黑", 30, INK, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub)
    set_cn_font(r2, "微软雅黑", 13, MUTED)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"（{date.today().strftime('%Y年%m月%d日')}）")
    set_cn_font(r3, "微软雅黑", 10, GOLD)

    doc.add_page_break()


def build_item(doc: Document, item: dict, idx: int):
    title = item.get("title", "")
    score = item.get("score", 0)
    cat = CAT_LABEL.get(item.get("category", ""), item.get("category", ""))
    kw = item.get("keyword", "")
    url = item.get("url", "")

    # 标题行
    p = doc.add_paragraph()
    r = p.add_run(f"{idx}. {title}")
    set_cn_font(r, "微软雅黑", 12, INK, bold=True)

    # 元信息行
    meta = []
    if score:
        meta.append(SCORE_LABEL.get(score, ""))
    if cat:
        meta.append(cat)
    if kw:
        meta.append(f"关键词：{kw}")
    if meta:
        pm = doc.add_paragraph()
        rm = pm.add_run("  |  ".join(meta))
        set_cn_font(rm, "微软雅黑", 8.5, RED if score >= 4 else MUTED)

    # 摘要
    summary = item.get("summary", "")
    if summary:
        ps = doc.add_paragraph()
        rs = ps.add_run(summary)
        set_cn_font(rs, "微软雅黑", 10, MUTED)

    # 数据点
    dp = item.get("data_points", "")
    if dp:
        pd = doc.add_paragraph()
        rd = pd.add_run(f"关键数据：{dp}")
        set_cn_font(rd, "微软雅黑", 9, GOLD)

    # 面试价值
    iv = item.get("interview_value", "")
    if iv:
        pi = doc.add_paragraph()
        ri = pi.add_run(f"💡 面试价值：{iv}")
        set_cn_font(ri, "微软雅黑", 9, RED)

    # 话术
    tip = item.get("interview_tip", "")
    if tip:
        pt = doc.add_paragraph()
        rt = pt.add_run(f"🗣 话术：{tip}")
        set_cn_font(rt, "微软雅黑", 9, GOLD)

    # 链接
    if url:
        pu = doc.add_paragraph()
        ru = pu.add_run(url)
        set_cn_font(ru, "微软雅黑", 8, MUTED)

    doc.add_paragraph()  # 间隔


def build_summary_table(doc: Document, items: list[dict]):
    """首页概览表：评分排序的简明清单"""
    p = doc.add_paragraph()
    r = p.add_run("一、本次动态概览")
    set_cn_font(r, "微软雅黑", 14, RED, bold=True)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["评分", "类别", "标题", "关键词"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        rh = cell.paragraphs[0].add_run(h)
        set_cn_font(rh, "微软雅黑", 9, INK, bold=True)

    for it in sorted(items, key=lambda x: -x.get("score", 0)):
        row = table.add_row()
        vals = [
            SCORE_LABEL.get(it.get("score", 0), ""),
            CAT_LABEL.get(it.get("category", ""), it.get("category", "")),
            it.get("title", ""),
            it.get("keyword", ""),
        ]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            rr = cell.paragraphs[0].add_run(v)
            set_cn_font(rr, "微软雅黑", 8.5, RED if i == 0 and it.get("score", 0) >= 4 else INK)

    doc.add_paragraph()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--date", default=date.today().isoformat())
    parser.add_argument("--out", default="")
    args = parser.parse_args()

    src = ANALYZED_DIR / args.date / f"analyzed_{args.date}.json"
    if not src.exists():
        print(f"错误：{src} 不存在。先运行 python analyze.py --date {args.date}")
        return 1
    data = json.loads(src.read_text(encoding="utf-8"))
    items = data.get("items", [])

    doc = Document()
    # 页面边距
    for section in doc.sections:
        section.left_margin = section.right_margin = Cm(2.2)
        section.top_margin = section.bottom_margin = Cm(2.2)

    build_cover(doc, "玄机科技 IP 动态速报", "国漫IP运营岗 · 面试备战素材")
    build_summary_table(doc, items)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("二、详细分析")
    set_cn_font(r, "微软雅黑", 14, RED, bold=True)

    for idx, it in enumerate(items, 1):
        build_item(doc, it, idx)

    out_path = Path(args.out) if args.out else OUTPUT_DIR / f"玄机IP动态速报_{args.date}.docx"
    doc.save(str(out_path))
    print(f"docx 生成成功：{out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
