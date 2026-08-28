"""
玄策 · 数据导入中心（Plan C：全链路智能导入）

流程：上传原始内容（CSV/JSON/JSONL/TXT）→ GLM-4-Flash 分批整理（清洗/去重/分类/情感/角色归属/摘要）
     → 预览确认 → 分批写入目标表（community_feedback / xuanji_feed / character_daily_metrics）。

GLM-4-Flash 能力：免费、支持 JSON 输出、可分批并发调用。
本项目已配置 zhipu_config.json（model=glm-4-flash + apiKey），复用 ai.llm 统一接入。
无 key / 调用失败时降级为规则整理（_rule_analyze），保证管道始终可用。
"""
from __future__ import annotations

import csv
import io
import json
import re
from datetime import date
from pathlib import Path

from database import get_db

# 目标表 → 允许字段
TARGETS = {
    "community_feedback": ["platform", "user_name", "content", "sentiment", "role_type", "date"],
    "xuanji_feed": ["title", "url", "summary", "category", "score", "interview_value", "keyword"],
    "character_daily_metrics": ["character_id", "date", "discussions", "search_index", "fan_growth", "fanworks", "commercial_score"],
    "metrics": ["platform", "followers", "reads_views", "interactions", "recorded_at"],
    "follower_history": ["date", "platform", "followers"],
}

# 结构化数值目标：无需 GLM 文本整理，上传即可提交（跳过 analyze）
NO_ANALYZE_TARGETS = {"metrics", "follower_history", "character_daily_metrics"}

_POS = ("好", "爱", "绝", "赞", "喜欢", "期待", "好看", "支持", "燃", "神", "美", "强", "牛", "推荐", "顶")
_NEG = ("差", "难看", "弃", "失望", "烂", "退", "烦", "太慢", "拉", "崩", "水", "糊", "骗")
_CHAR_NAMES = ("盖聂", "天明", "少司命", "卫庄", "雪女", "韩非", "焰灵姬", "紫女",
               "武庚", "白菜", "唐三", "小舞", "比比东", "李景珑", "孔鸿俊",
               "罗峰", "李长寿", "霍雨浩", "唐舞桐")
_IP_KEYWORDS = {
    "秦时明月": ("秦时明月", "秦时", "明月", "百步飞剑", "沧海横流", "盖聂", "天明", "少司命", "卫庄", "雪女"),
    "天行九歌": ("天行九歌", "天行", "韩非", "焰灵姬", "紫女"),
    "武庚纪": ("武庚纪", "武庚", "白菜"),
    "斗罗大陆": ("斗罗大陆", "斗罗", "唐三", "小舞", "比比东", "绝世唐门", "霍雨浩", "唐舞桐"),
    "天宝伏妖录": ("天宝伏妖录", "天宝", "李景珑", "孔鸿俊"),
    "吞噬星空": ("吞噬星空", "吞噬", "罗峰"),
    "师兄啊师兄": ("师兄啊师兄", "师兄", "李长寿"),
    "牧神记": ("牧神记",),
    "天谕": ("天谕",),
}

SYSTEM_PROMPT = """你是玄机科技（国漫IP运营）的数据整理助手。把用户输入的原始文本逐条整理成结构化 JSON 数组，每项字段：
- content: 清洗后的正文（去多余空白、截断 300 字）
- sentiment: positive / neutral / negative（情感）
- role_type: 路人 / 角色党 / 美术党 / 剧情党（用户类型）
- ip_name: 命中哪个玄机 IP（秦时明月/天行九歌/武庚纪/斗罗大陆/天宝伏妖录/吞噬星空/师兄啊师兄/牧神记/天谕），未命中填 ""
- character_name: 命中哪个角色（盖聂/天明/少司命/卫庄/雪女/韩非/焰灵姬/紫女/唐三/小舞/比比东/武庚/白菜/李景珑/孔鸿俊/罗峰/李长寿/霍雨浩/唐舞桐），未命中填 ""
- summary: 20 字内摘要
只输出 JSON 数组，不要输出任何其他文字。"""


def _parse_file(filename: str, raw: bytes) -> list[dict]:
    """按扩展名解析上传文件 → 统一行列表 [{"content": str, "extra": {...}}]"""
    name = (filename or "").lower()
    text = raw.decode("utf-8", errors="replace")
    rows: list[dict] = []
    if name.endswith(".jsonl") or name.endswith(".ndjson"):
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                rows.append(_pick_text(obj))
            except Exception:
                rows.append({"content": line})
    elif name.endswith(".json"):
        try:
            data = json.loads(text)
            if isinstance(data, list):
                for obj in data:
                    rows.append(_pick_text(obj))
            elif isinstance(data, dict):
                rows.append(_pick_text(data))
        except Exception:
            rows.append({"content": text})
    elif name.endswith(".csv"):
        reader = csv.DictReader(io.StringIO(text))
        for row in reader:
            content = row.get("content") or row.get("正文") or row.get("text") or ""
            rows.append({"content": str(content).strip(), **row})
    else:  # txt / 其他
        for line in text.splitlines():
            line = line.strip()
            if line:
                rows.append({"content": line})
    return rows


def _pick_text(obj: dict) -> dict:
    content = (obj.get("content") or obj.get("desc") or obj.get("title")
               or obj.get("正文") or obj.get("text") or "")
    return {"content": str(content).strip(), **{k: v for k, v in obj.items() if k != "content"}}


def _rule_analyze(content: str) -> dict:
    """无 LLM key 降级：规则情感 + 角色/IP 归属（与 community._classify 同源）"""
    content = content.strip()[:300]
    if any(w in content for w in _NEG):
        sentiment = "negative"
    elif any(w in content for w in _POS):
        sentiment = "positive"
    else:
        sentiment = "neutral"
    role = "路人"
    if any(n in content for n in _CHAR_NAMES):
        role = "角色党"
    elif any(k in content for k in ("美术", "画风", "壁纸", "立绘", "原画")):
        role = "美术党"
    elif any(k in content for k in ("剧情", "世界观", "设定", "结局", "主线")):
        role = "剧情党"
    ip_name, char_name = "", ""
    for ipn, kws in _IP_KEYWORDS.items():
        if any(k in content for k in kws):
            ip_name = ipn
            for c in _CHAR_NAMES:
                if c in content:
                    char_name = c
                    break
            break
    return {
        "content": content, "sentiment": sentiment, "role_type": role,
        "ip_name": ip_name, "character_name": char_name,
        "summary": content[:20],
    }


def analyze_batch(items: list[dict], use_llm: bool = True, batch_size: int = 0) -> list[dict]:
    """分批整理：LLM（glm-4-flash，JSON 输出）或规则降级。返回结构化行列表。
    batch_size<=0 时自动缩放：>200 条用 50/批（减少调用次数），否则 20/批。"""
    from ai.llm import resolve_llm, llm_chat
    results: list[dict] = []
    cfg = resolve_llm()
    can_llm = use_llm and cfg is not None
    if batch_size <= 0:
        batch_size = 50 if len(items) > 200 else 20

    for start in range(0, len(items), batch_size):
        chunk = items[start:start + batch_size]
        texts = [r.get("content", "") for r in chunk]
        if can_llm:
            try:
                user_prompt = json.dumps(texts, ensure_ascii=False)
                content = llm_chat([
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": user_prompt},
                ], temperature=0.2, timeout=90)
                # 容错解析：去掉 markdown 围栏 / 前缀文字
                m = re.search(r"\[[\s\S]*\]", content)
                parsed = json.loads(m.group(0)) if m else []
                if not isinstance(parsed, list):
                    parsed = []
                for i, row in enumerate(parsed[:len(chunk)]):
                    if isinstance(row, dict) and row.get("content"):
                        results.append(row)
                    else:
                        results.append(_rule_analyze(texts[i] if i < len(texts) else ""))
                # 批次缺行补齐
                if len(results) < len(chunk):
                    for i in range(len(results), len(chunk)):
                        results.append(_rule_analyze(texts[i] if i < len(texts) else ""))
                continue
            except Exception as e:
                print(f"[import] LLM batch failed ({start//batch_size}), fallback rules: {e}")
        for t in texts:
            results.append(_rule_analyze(t))
    return results


def _dedup_existing(cur, target: str) -> set[str]:
    """目标表已存在内容 hash（用于导入去重）"""
    seen: set[str] = set()
    if target == "community_feedback":
        for r in cur.execute("SELECT content, platform, user_name FROM community_feedback"):
            seen.add(hashlib_md5(f"{r['content']}|{r['platform']}|{r['user_name']}"))
    elif target == "xuanji_feed":
        for r in cur.execute("SELECT title FROM xuanji_feed"):
            seen.add(hashlib_md5(f"title|{r['title']}"))
    return seen


def hashlib_md5(s: str) -> str:
    import hashlib
    return hashlib.md5(s.encode("utf-8")).hexdigest()


def write_batch(cur, task_id: int, seq: int, rows: list[dict], target: str, seen: set[str]) -> dict:
    """写入一批到目标表（去重 + 来源标记 source='import'）。返回 {inserted, skipped, message}"""
    inserted = skipped = 0
    today = date.today().isoformat()

    if target == "community_feedback":
        for row in rows:
            content = str(row.get("content") or "").strip()[:300]
            if len(content) < 4:
                continue
            platform = str(row.get("platform") or "其他")[:20]
            uname = str(row.get("user_name") or "导入用户")[:20]
            key = hashlib_md5(f"{content}|{platform}|{uname}")
            if key in seen:
                skipped += 1
                continue
            cur.execute(
                """INSERT INTO community_feedback
                   (platform,user_name,content,sentiment,role_type,date,source,crawled_at)
                   VALUES (?,?,?,?,?,?,?,?)""",
                (platform, uname, content,
                 row.get("sentiment") or "neutral", row.get("role_type") or "路人",
                 row.get("date") or today, f"import:{task_id}", today))
            seen.add(key)
            inserted += 1

    elif target == "xuanji_feed":
        for row in rows:
            title = str(row.get("title") or row.get("content") or "").strip()[:120]
            if not title:
                continue
            key = hashlib_md5(f"title|{title}")
            if key in seen:
                skipped += 1
                continue
            cur.execute(
                """INSERT INTO xuanji_feed (fetch_date,keyword,category,title,url,summary,score,interview_value,raw_content)
                   VALUES (?,?,?,?,?,?,?,?,?)""",
                (row.get("date") or today, row.get("keyword") or "导入", row.get("category") or "ip",
                 title, str(row.get("url") or "")[:500], str(row.get("summary") or "")[:300],
                 int(row.get("score") or 0), str(row.get("interview_value") or "")[:200],
                 json.dumps({**row, "source": f"import:{task_id}"}, ensure_ascii=False)[:2000]))
            seen.add(key)
            inserted += 1

    elif target == "metrics":
        # 平台快照导入：CSV 每行一个平台（platform/followers/reads_views/interactions/recorded_at）
        for row in rows:
            platform = str(row.get("platform") or row.get("平台") or "").strip()
            if not platform:
                continue
            pkey = platform
            if pkey in ("B站", "b站"):
                pkey = "bilibili"
            elif pkey in ("微博",):
                pkey = "weibo"
            elif pkey in ("小红书",):
                pkey = "xiaohongshu"
            elif pkey in ("公众号", "微信"):
                pkey = "wechat"
            rec_at = str(row.get("recorded_at") or row.get("日期") or today)
            cur.execute("SELECT id FROM metrics WHERE platform=? AND recorded_at=?", (pkey, rec_at))
            dup = cur.fetchone()
            if dup:
                skipped += 1
                continue
            # followers 缺省(0) 时沿用该平台已有最新粉丝，避免覆盖为 0
            followers = int(row.get("followers") or row.get("粉丝") or 0)
            if followers <= 0:
                cur.execute(
                    "SELECT followers FROM metrics WHERE platform=? ORDER BY recorded_at DESC LIMIT 1",
                    (pkey,))
                prev = cur.fetchone()
                if prev and prev["followers"]:
                    followers = prev["followers"]
            cur.execute(
                """INSERT INTO metrics (platform,followers,reads_views,interactions,engagement_rate,top_content,recorded_at,source)
                   VALUES (?,?,?,?,?,?,?,?)""",
                (pkey, followers,
                 int(row.get("reads_views") or row.get("播放") or row.get("阅读") or 0),
                 int(row.get("interactions") or row.get("互动") or 0),
                 float(row.get("engagement_rate") or 0), "[]", rec_at, f"import:{task_id}"))
            inserted += 1

    elif target == "follower_history":
        # 粉丝历史导入：CSV 每行 date/platform/followers
        for row in rows:
            d = str(row.get("date") or row.get("日期") or "").strip()
            platform = str(row.get("platform") or row.get("平台") or "").strip()
            if not d or not platform:
                continue
            pkey = platform
            if pkey in ("B站", "b站"):
                pkey = "bilibili"
            elif pkey in ("微博",):
                pkey = "weibo"
            elif pkey in ("小红书",):
                pkey = "xiaohongshu"
            elif pkey in ("公众号", "微信"):
                pkey = "wechat"
            cur.execute("SELECT id FROM follower_history WHERE date=? AND platform=?", (d, pkey))
            dup = cur.fetchone()
            if dup:
                skipped += 1
                continue
            cur.execute(
                "INSERT INTO follower_history (date,platform,followers,source) VALUES (?,?,?,?)",
                (d, pkey, int(row.get("followers") or row.get("粉丝") or 0), f"import:{task_id}"))
            inserted += 1

    elif target == "character_daily_metrics":
        for row in rows:
            cid = row.get("character_id")
            d = str(row.get("date") or today)
            if not cid or not d:
                continue
            cur.execute("SELECT id FROM characters WHERE id=?", (cid,))
            if not cur.fetchone():
                continue
            cur.execute(
                "SELECT id FROM character_daily_metrics WHERE character_id=? AND date=? AND source LIKE 'import%'",
                (cid, d))
            dup = cur.fetchone()
            if dup:
                skipped += 1
                continue
            cur.execute(
                """INSERT INTO character_daily_metrics
                   (character_id,date,search_index,discussions,fan_growth,fanworks,commercial_score,source)
                   VALUES (?,?,?,?,?,?,?,?)""",
                (cid, d, int(row.get("search_index") or 0), int(row.get("discussions") or 0),
                 int(row.get("fan_growth") or 0), int(row.get("fanworks") or 0),
                 float(row.get("commercial_score") or 0), f"import:{task_id}"))
            inserted += 1

    return {"inserted": inserted, "skipped": skipped, "message": f"批次{seq}: +{inserted} / 去重跳过 {skipped}"}


def create_task(name: str, source_type: str, target: str, rows: list[dict], model: str) -> int:
    conn = get_db(); cur = conn.cursor()
    cur.execute(
        """INSERT INTO import_tasks (name,source_type,target,status,total,model,payload)
           VALUES (?,?,?,?,?,?,?)""",
        (name, source_type, target, "pending", len(rows), model, json.dumps(rows, ensure_ascii=False)))
    tid = cur.lastrowid
    conn.commit(); conn.close()
    return tid


def get_task(tid: int) -> dict | None:
    conn = get_db(); cur = conn.cursor()
    cur.execute("SELECT * FROM import_tasks WHERE id=?", (tid,))
    row = cur.fetchone()
    conn.close()
    if not row:
        return None
    d = dict(row)
    d["payload"] = json.loads(d.get("payload") or "[]")
    d["errors"] = json.loads(d.get("errors") or "[]")
    return d


def update_task(tid: int, **kw) -> None:
    conn = get_db(); cur = conn.cursor()
    cols = ", ".join(f"{k}=?" for k in kw)
    cur.execute(f"UPDATE import_tasks SET {cols} WHERE id=?", (*kw.values(), tid))
    conn.commit(); conn.close()


def list_tasks(limit: int = 50) -> list[dict]:
    conn = get_db(); cur = conn.cursor()
    cur.execute("SELECT * FROM import_tasks ORDER BY id DESC LIMIT ?", (limit,))
    rows = [dict(r) for r in cur.fetchall()]
    conn.close()
    for r in rows:
        r["errors"] = json.loads(r.get("errors") or "[]")
    return rows


def log_batch(cur, task_id: int, seq: int, result: dict) -> None:
    cur.execute(
        """INSERT INTO import_batches (task_id,seq,status,inserted,skipped,message)
           VALUES (?,?,?,?,?,?)""",
        (task_id, seq, "done", result.get("inserted", 0), result.get("skipped", 0), result.get("message", "")))


# ==================== 速报 docx 结构化解析（数据SOP：智普agent 抓取 → 导入中心载入） ====================
# 《玄机IP动态速报》docx 本身已是结构化简报（标题层级 + 【标签】条目 + 话术表格），
# 规则解析比 GLM 通用整理更可靠且零成本。映射约定：
#   二、各IP动态详情   → category=ip（标签/小节自动归属真实 IP）
#   三、玄机公司动态   → category=ipo（IPO 子题）/ company
#   四、面试素材表格   → category=strategy，话术全文进 summary 与 interview_value
#   五、本周趋势回顾   → category=strategy 周回顾单条
#   一、要闻速览为详情区索引，跳过以免入库重复。
# 期号日期优先取文件名（玄机IP动态速报_YYYY-MM-DD.docx），回退正文首段，保证补录历史期号时间正确。

_L1_RE = re.compile(r"^[一二三四五六七八九十]+\s*[、.]")
_L2_RE = re.compile(r"^\d{1,2}\s*[.、．]\s*(.+)$")
_TAG_RE = re.compile(r"^【([^】]{1,40})】\s*(.*)$", re.S)
_BULLET_LABELED_RE = re.compile(r"^[•·\-–—]\s*([^：:【】]{2,24})[：:]\s*(.+)$", re.S)
_DATE_IN_NAME_RE = re.compile(r"(20\d{2})[-._]?(\d{2})[-._]?(\d{2})")
_DATE_IN_TEXT_RE = re.compile(r"(20\d{2})[-年/.](\d{1,2})[-月/.](\d{1,2})")
_HOT_WORDS = ("确认", "验证", "判据", "溯源", "沉淀", "提案", "观察点")


def _match_ip(text: str) -> str:
    if "绝世唐门" in text:
        return "斗罗大陆"
    for ipn in _IP_KEYWORDS:
        if ipn in text:
            return ipn
    return ""


def _docx_texts(raw: bytes) -> tuple[list[str], list[list[str]]]:
    """docx 字节 → (非空段落文本列表, 表格行文本列表)。python-docx 与 pipeline/generate_docx.py 同源依赖。"""
    from docx import Document  # 延迟导入：仅速报 docx 载入时需要
    doc = Document(io.BytesIO(raw))
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    table_rows: list[list[str]] = []
    for tb in doc.tables:
        for row in tb.rows:
            cells = [" ".join(c.text.split()) for c in row.cells]
            if any(cells):
                table_rows.append(cells)
    return paras, table_rows


def _bulletin_date(filename: str, paras: list[str]) -> str:
    m = _DATE_IN_NAME_RE.search(filename or "")
    if not m:
        for t in paras[:12]:
            m = _DATE_IN_TEXT_RE.search(t)
            if m:
                break
    if m:
        try:
            return f"{int(m.group(1)):04d}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
        except ValueError:
            pass
    return date.today().isoformat()


def _clean_subject(subject: str) -> str:
    return (subject or "").split("：")[0].split(":")[0].strip() or "本期动态"


def _feed_row(label: str, tag: str, body: str, bdate: str, category: str) -> dict:
    score = 5 if "重点" in tag else (4 if any(w in tag for w in _HOT_WORDS) else 3)
    title = label if tag == label else f"{label}｜{tag}"
    return {
        "title": title[:120], "url": "", "summary": body[:300],
        "category": category, "score": score, "interview_value": "",
        "keyword": label, "date": bdate,
    }


def parse_bulletin_docx(filename: str, raw: bytes) -> list[dict]:
    """《玄机IP动态速报》docx → xuanji_feed 结构化行列表（规则解析，零 LLM 依赖）。"""
    paras, table_rows = _docx_texts(raw)
    bdate = _bulletin_date(filename, paras)
    rows: list[dict] = []
    weekly: list[str] = []

    section = ""        # 当前一级章节（一/二/三/四/五…）
    subject = ""        # 当前二级小节（IP 名 / 公司子题）
    tag, buf = "", []   # 当前【标签】块

    def flush() -> None:
        nonlocal tag, buf
        if tag and buf:
            body = "\n".join(buf).strip()
            if section == "二":
                subject_clean = _clean_subject(subject)
                label = _match_ip(tag) or _match_ip(subject_clean) or subject_clean
                rows.append(_feed_row(label, tag, body, bdate, "ip"))
            elif section == "三":
                is_ipo = "IPO" in subject.upper() or "IPO" in tag.upper()
                row = _feed_row("玄机IPO" if is_ipo else "玄机公司", tag, body, bdate,
                                "ipo" if is_ipo else "company")
                row["keyword"] = "玄机科技"
                rows.append(row)
        tag, buf = "", []

    for text in paras:
        if _L1_RE.match(text):
            flush()
            section = text[0]
            subject = ""
            continue
        if section == "五":  # 周五回顾章节：含【标签】行在内全部并入周回顾
            weekly.append(text)
            continue
        m2 = _L2_RE.match(text)
        if m2 and section in ("二", "三"):
            flush()
            subject = m2.group(1).strip()
            continue
        mt = _TAG_RE.match(text)
        if mt:
            flush()
            tag, buf = mt.group(1).strip(), [mt.group(2).strip()]
            continue
        if section in ("二", "三"):
            mb = _BULLET_LABELED_RE.match(text)
            if mb:  # “• 标签：内容”行 → 独立成块
                flush()
                tag, buf = mb.group(1).strip()[:40], [mb.group(2).strip()]
            elif tag:
                buf.append(text)
            else:  # 小节内无标签散行 → 以小节主题兜底成块
                tag, buf = _clean_subject(subject)[:40], [text]
    flush()

    # 四、面试素材（表格）：话术全文 → summary，前 200 字同时进 interview_value（raw_content 保留全量）
    for cells in table_rows:
        text = " ".join(c for c in cells if c).strip()
        if "话术" not in text:
            continue
        head = next((c for c in cells if c), "")
        title = head.split("「")[0].strip()[:120] or "面试素材"
        rows.append({
            "title": title, "url": "", "summary": text[:300],
            "category": "strategy", "score": 4,
            "interview_value": text[:200], "keyword": "面试素材", "date": bdate,
        })

    # 五、本周趋势回顾 → 单条周回顾行（周五特供）
    if weekly:
        rows.append({
            "title": f"玄机公司｜本周趋势回顾（{bdate}）", "url": "",
            "summary": "\n".join(weekly)[:300], "category": "strategy", "score": 4,
            "interview_value": "", "keyword": "玄机科技", "date": bdate,
        })
    return rows
